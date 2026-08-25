from docx import Document
from docx.shared import Pt

PATH = r"建模总控/03-deliverables/02_Q2_碰撞模型/02_Q2_碰撞模型_论文手审阅.docx"

doc = Document(PATH)
for p in doc.paragraphs:
    if p.text.startswith("使用相同的实体矩形和 SAT 代码，分别以 1 s、2 s 步长"):
        p.text = (
            "先以 0.5 s、1 s、2 s 步长复核根精化稳定性；三种步长均回到同一临界时间。"
            "但这项结果本身不能排除粗采样点之间的短暂负裕度，因此另行执行 0.1 s 完整扫描和局部极小值审计。"
        )
    if p.text.startswith("扫描还发现 436 s 以后存在再次的非单调"):
        p.text = "扫描还发现 436 s 以后存在再次的非单调碰撞转变，因此不能把“首次碰撞后一直碰撞”当作模型假设。问题2只报告首次达到碰撞临界的终止时刻。"

# Make the edit idempotent by removing a previously inserted audit block.
paragraphs = list(doc.paragraphs)
start = next((i for i, p in enumerate(paragraphs) if p.text == "5.1 早期窄碰撞漏检审计"), None)
end = next(i for i, p in enumerate(paragraphs) if p.text.startswith("6. 论文表达与局限"))
if start is not None:
    for p in paragraphs[start:end]:
        p._element.getparent().remove(p._element)

target = next(p for p in doc.paragraphs if p.text.startswith("6. 论文表达与局限"))
target.insert_paragraph_before("5.1 早期窄碰撞漏检审计", style="Heading 2")
target.insert_paragraph_before(
    "仅比较 0.5 s、1 s 和 2 s 的根精化结果，不能单独排除采样点之间持续时间很短的负裕度区间。"
    "为此，对 0≤t≤412.5 s 做完整 0.1 s 扫描，共 4,126 个时刻。扫描只发现一个正裕度到非正裕度的区间：[412.4, 412.5] s；"
    "其 Brent 精化时间为 412.4738376822 s，与主结果 412.4738376821 s 一致。"
)
target.insert_paragraph_before(
    "在 0≤t≤412 s 范围内，4,121 个采样点的全局裕度全部为正。最小采样裕度为 8.73218×10⁻⁴ m，出现在 409.8 s。"
    "共识别 67 个采样局部极小点，并在各自左右相邻 0.1 s 窗口内连续精化；最小精化裕度为 8.65235×10⁻⁴ m，位于 409.8319894 s，"
    "其左右采样裕度分别为 1.000806×10⁻³ m 和 9.014010×10⁻⁴ m，仍明显高于零。"
)
target.insert_paragraph_before(
    "为消除采样相位影响，使用 0.05 s 错位网格复核，最小裕度为 8.67767×10⁻⁴ m；再对 410--412.5 s 以 0.001 s 加密扫描，"
    "唯一符号跨越为 [412.473, 412.474] s。因此，早于 412.4738376821 s 的窄负裕度区间没有运行证据支持。"
)

doc.save(PATH)
